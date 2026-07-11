"""Parse participant names from Word/RTF files."""

from __future__ import annotations

import re
from pathlib import Path

from .common import ResultRow, make_row
from .text_parser import parse_text_file


def parse_office_file(path: Path, event_id: int) -> list[ResultRow]:
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return _parse_docx(path, event_id)
    if suffix == ".rtf":
        return _parse_rtf(path, event_id)
    if suffix == ".doc":
        return _parse_doc(path, event_id)
    return []


def _parse_docx(path: Path, event_id: int) -> list[ResultRow]:
    from docx import Document

    document = Document(path)
    lines = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            lines.append("\t".join(cell.text.strip() for cell in row.cells))
    temp = path.with_suffix(".extracted.txt")
    temp.write_text("\n".join(lines), encoding="utf-8")
    try:
        return parse_text_file(temp, event_id)
    finally:
        temp.unlink(missing_ok=True)


def _parse_rtf(path: Path, event_id: int) -> list[ResultRow]:
    from striprtf.striprtf import rtf_to_text

    text = rtf_to_text(path.read_text(encoding="utf-8", errors="replace"))
    temp = path.with_suffix(".extracted.txt")
    temp.write_text(text, encoding="utf-8")
    try:
        return parse_text_file(temp, event_id)
    finally:
        temp.unlink(missing_ok=True)


def _parse_doc(path: Path, event_id: int) -> list[ResultRow]:
    # Legacy .doc: extract printable strings as fallback
    raw = path.read_bytes()
    chunks = re.findall(rb"[\x20-\x7e\xc0-\xff]{4,}", raw)
    text = "\n".join(chunk.decode("cp1252", errors="ignore") for chunk in chunks)
    temp = path.with_suffix(".extracted.txt")
    temp.write_text(text, encoding="utf-8")
    try:
        parsed = parse_text_file(temp, event_id)
        for row in parsed:
            row.parse_source = "doc"
            row.parse_confidence = "low"
        return parsed
    finally:
        temp.unlink(missing_ok=True)
